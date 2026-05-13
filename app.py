import streamlit as st
from docx import Document
from fpdf import FPDF
from io import BytesIO
from datetime import datetime
import difflib
import re
import json
import fitz  # pymupdf
from dotenv import load_dotenv
import os
from google import genai

load_dotenv()

client_gemini = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))

def gemini(prompt):
    response = client_gemini.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt
    )
    return response.text

# --- SESSION STATE INITIALISIEREN ---
if "anzeige" not in st.session_state:
    st.session_state.anzeige = None
if "historie" not in st.session_state:
    st.session_state.historie = []
if "hr_status" not in st.session_state:
    st.session_state.hr_status = None
if "hr_kommentar" not in st.session_state:
    st.session_state.hr_kommentar = ""
if "job_titel_gespeichert" not in st.session_state:
    st.session_state.job_titel_gespeichert = ""
if "compliance" not in st.session_state:
    st.session_state.compliance = None
if "reset_trigger" not in st.session_state:
    st.session_state.reset_trigger = 0
if "aktueller_schritt" not in st.session_state:
    st.session_state.aktueller_schritt = 1
if "agent_freitext" not in st.session_state:
    st.session_state.agent_freitext = ""
if "agent_rückfragen" not in st.session_state:
    st.session_state.agent_rückfragen = None
if "agent_antworten" not in st.session_state:
    st.session_state.agent_antworten = ""

# --- HR EINSTELLUNGEN SPEICHERN/LADEN ---
def lade_hr_einstellungen():
    try:
        with open("hr_einstellungen.json", "r", encoding="utf-8") as f:
            return json.load(f)
    except:
        return {
            "firmen_name": "Alwina Digital",
            "anrede": "Du",
            "gender_stil": "Gender-Sternchen (z.B. Mitarbeiter*innen)",
            "highlight": "Dachterrasse in Wiesbaden",
            "stufen_liste": [
                {"name": "Junior", "erfahrung_min": 0, "erfahrung_max": 2, "gehalt_min": 35000, "gehalt_max": 50000},
                {"name": "Professional", "erfahrung_min": 2, "erfahrung_max": 5, "gehalt_min": 50000, "gehalt_max": 70000},
                {"name": "Senior", "erfahrung_min": 5, "erfahrung_max": 10, "gehalt_min": 70000, "gehalt_max": 90000},
                {"name": "Lead", "erfahrung_min": 10, "erfahrung_max": 99, "gehalt_min": 90000, "gehalt_max": 150000},
            ]
        }

def speichere_hr_einstellungen(einstellungen):
    with open("hr_einstellungen.json", "w", encoding="utf-8") as f:
        json.dump(einstellungen, f, ensure_ascii=False, indent=2)

# --- HILFSFUNKTIONEN ---
def formatiere_ueberschriften(text):
    zeilen = text.split("\n")
    ergebnis = []
    for zeile in zeilen:
        if zeile.startswith("# "):
            ergebnis.append("### " + zeile.lstrip("#").strip().upper())
        elif zeile.startswith("## "):
            ergebnis.append("#### " + zeile.lstrip("#").strip().upper())
        else:
            ergebnis.append(zeile)
    return "\n".join(ergebnis)

def hole_wissen_aus_archiv(job_titel):
    archiv_text = ""
    archiv_ordner = "archiv"
    if os.path.exists(archiv_ordner):
        for datei in os.listdir(archiv_ordner):
            if datei.endswith(".pdf"):
                try:
                    doc = fitz.open(os.path.join(archiv_ordner, datei))
                    for seite in doc:
                        archiv_text += seite.get_text()
                    archiv_text += "\n---\n"
                except:
                    pass
    if not archiv_text:
        try:
            with open("archiv.txt", "r", encoding="utf-8") as f:
                archiv_text = f.read()
        except:
            return "Nutze einen modernen IT-Stil."
    abschnitte = archiv_text.split("---")
    treffer = [a.strip() for a in abschnitte if any(
        wort.lower() in a.lower() for wort in job_titel.split()
    )]
    if treffer:
        return "\n---\n".join(treffer[:2])
    elif abschnitte:
        return abschnitte[0].strip()
    else:
        return "Nutze einen modernen IT-Stil."

def fuehre_compliance_loop_durch(anzeige_text):
    aktuelle_anzeige = anzeige_text
    for versuch in range(1, 4):
        compliance_prompt = f"""
        Du bist ein Experte für deutsches Arbeitsrecht und diskriminierungsfreie Sprache.
        Analysiere diese Stellenanzeige auf problematische Formulierungen:
        {aktuelle_anzeige}
        Prüfe auf: Altersdiskriminierung, Geschlechterdiskriminierung, Sprachdiskriminierung,
        Überforderungs-Begriffe, Unrealistische Anforderungen, Versteckte Diskriminierung.
        Antworte NUR in diesem Format:
        STATUS: [BESTANDEN oder PROBLEME GEFUNDEN]
        PROBLEME:
        - [Problem 1]
        EMPFEHLUNG:
        - [Verbesserung 1]
        Falls keine Probleme: schreibe bei PROBLEME: "Keine gefunden ✅"
        """
        ergebnis = gemini(compliance_prompt)
        if "BESTANDEN" in ergebnis:
            return aktuelle_anzeige
        if versuch < 3:
            korrektur_prompt = f"""
            Diese Stellenanzeige hat Compliance-Probleme:
            {ergebnis}
            Aktuelle Anzeige:
            {aktuelle_anzeige}
            Korrigiere NUR die problematischen Stellen.
            Schreibe die vollständige korrigierte Anzeige zurück.
            """
            aktuelle_anzeige = formatiere_ueberschriften(gemini(korrektur_prompt))
    return aktuelle_anzeige

def generiere_anzeige(job_titel, level, erfahrung, ort, homeoffice, aufgaben):
    aktuelle_daten = {
        "titel": job_titel,
        "level": level,
        "erfahrung": erfahrung,
        "ort": ort,
        "homeoffice": homeoffice
    }
    ergebnisse = pruefe_ausschreibung(aktuelle_daten)
    if ergebnisse:
        return None, ergebnisse

    st.session_state.job_titel_gespeichert = job_titel
    relevantes_wissen = hole_wissen_aus_archiv(job_titel)
    system_prompt = f"""
    Du bist ein HR-Experte für {FIRMEN_NAME}.
    
    WICHTIG: Der Firmenname ist IMMER "{FIRMEN_NAME}".
    Erwähne NIEMALS andere Firmennamen wie "Schwarz Digits", "Lidl", "Kaufland"
    oder andere Unternehmen aus den Vorlagen.
    Die Vorlagen dienen NUR als Stil-Referenz – nicht als Inhalt!
    
    STIL-VORLAGE AUS UNSEREM ARCHIV (Nutze NUR den Stil, nicht den Inhalt!):
    {relevantes_wissen}
    
    AUFTRAG:
    Schreibe eine neue, begeisternde Stellenanzeige für die Position: {job_titel}.
    Nutze die Tonalität und Struktur der Vorlage (Impact, Aufgaben, Profil).
    
    REGELN:
    - Anrede: {ANREDE}
    - Gender-Stil: {GENDER_STIL}
    - Highlight: {HIGHLIGHT}
    - Erfahrungslevel: {level}
    - Arbeitsmodell: {ort}
    
    GEHALT: Erwähne KEIN konkretes Gehalt.
    Nutze elegante Formulierungen wie: 'Dich erwartet ein attraktives Vergütungspaket'.
    
    INPUT-DETAILS:
    - Aufgaben: {aufgaben}
    
    Schreibe die Anzeige in schönem Markdown-Format.
    Alle Überschriften NUR in Großbuchstaben.
    """
    anzeige = formatiere_ueberschriften(gemini(system_prompt))
    return anzeige, []

def erstelle_word_dokument(anzeige_text):
    doc = Document()
    titel = doc.add_heading("Stellenanzeige", 0)
    titel.alignment = 1
    for zeile in anzeige_text.split("\n"):
        zeile = zeile.strip()
        if not zeile:
            doc.add_paragraph("")
        elif zeile.startswith("## "):
            doc.add_heading(zeile.replace("## ", ""), level=2)
        elif zeile.startswith("# "):
            doc.add_heading(zeile.replace("# ", ""), level=1)
        elif zeile.startswith("- "):
            doc.add_paragraph(zeile.replace("- ", ""), style="List Bullet")
        else:
            absatz = doc.add_paragraph()
            teile = re.split(r'(\*\*.*?\*\*)', zeile)
            for teil in teile:
                if teil.startswith("**") and teil.endswith("**"):
                    lauf = absatz.add_run(teil[2:-2])
                    lauf.bold = True
                else:
                    absatz.add_run(teil)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def erstelle_pdf_dokument(anzeige_text, job_titel):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    links = 15
    pdf.set_left_margin(links)
    nutzbare_breite = pdf.w - (2 * links)

    def clean(text):
        if not text: return ""
        t = text.replace('•', '-').replace('€', 'Euro').replace('–', '-')
        return t.encode('latin-1', 'replace').decode('latin-1')

    pdf.set_font("Helvetica", "B", 16)
    pdf.set_x(links)
    pdf.cell(nutzbare_breite, 10, clean(f"Stellenanzeige: {job_titel}"), ln=True, align="C")
    pdf.ln(5)
    for zeile in anzeige_text.split("\n"):
        zeile = zeile.strip()
        if not zeile:
            pdf.ln(4)
            continue
        pdf.set_x(links)
        if zeile.startswith("#"):
            saubere_zeile = clean(zeile.lstrip("#").strip())
            pdf.set_font("Helvetica", "B", 12)
            pdf.multi_cell(nutzbare_breite, 8, saubere_zeile)
            pdf.ln(2)
        elif zeile.startswith("- "):
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(nutzbare_breite, 6, clean(f"- {zeile[2:]}"))
        else:
            saubere_zeile = clean(re.sub(r'\*\*(.*?)\*\*', r'\1', zeile))
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(nutzbare_breite, 6, saubere_zeile)
    buffer = BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer

def zeige_diff_wortweise(alt_zeile, neu_zeile):
    alt_woerter = alt_zeile.split()
    neu_woerter = neu_zeile.split()
    matcher = difflib.SequenceMatcher(None, alt_woerter, neu_woerter)
    ergebnis = []
    for opcode, a0, a1, b0, b1 in matcher.get_opcodes():
        if opcode == "equal":
            ergebnis.append(" ".join(neu_woerter[b0:b1]))
        elif opcode == "insert":
            woerter = " ".join(neu_woerter[b0:b1])
            ergebnis.append(f"<span style='background-color:#00cc44;color:white;padding:2px 4px;border-radius:4px'>{woerter}</span>")
        elif opcode == "delete":
            woerter = " ".join(alt_woerter[a0:a1])
            ergebnis.append(f"<span style='background-color:#ff4444;color:white;padding:2px 4px;border-radius:4px;text-decoration:line-through'>{woerter}</span>")
        elif opcode == "replace":
            alt_w = " ".join(alt_woerter[a0:a1])
            neu_w = " ".join(neu_woerter[b0:b1])
            ergebnis.append(f"<span style='background-color:#ff4444;color:white;padding:2px 4px;border-radius:4px;text-decoration:line-through'>{alt_w}</span> <span style='background-color:#00cc44;color:white;padding:2px 4px;border-radius:4px'>{neu_w}</span>")
    return " ".join(ergebnis)

def zeige_schritt_anzeige(aktueller_schritt):
    schritte = ["📝 Daten eingeben", "✍️ Anzeige prüfen", "📤 An HR senden", "📥 Export"]
    cols = st.columns(4)
    for i, (col, schritt) in enumerate(zip(cols, schritte)):
        with col:
            if i + 1 < aktueller_schritt:
                st.markdown(f"<div style='text-align:center;padding:8px;background-color:#00cc4422;border-radius:8px;border:1px solid #00cc44'><small>✅ Schritt {i+1}</small><br><b style='font-size:12px'>{schritt}</b></div>", unsafe_allow_html=True)
            elif i + 1 == aktueller_schritt:
                st.markdown(f"<div style='text-align:center;padding:8px;background-color:#3b82f622;border-radius:8px;border:2px solid #3b82f6'><small>👉 Schritt {i+1}</small><br><b style='font-size:12px'>{schritt}</b></div>", unsafe_allow_html=True)
            else:
                st.markdown(f"<div style='text-align:center;padding:8px;background-color:#2a2a35;border-radius:8px;border:1px solid #3a3a45'><small style='color:#6b6b7a'>Schritt {i+1}</small><br><b style='font-size:12px;color:#6b6b7a'>{schritt}</b></div>", unsafe_allow_html=True)

def pruefe_ausschreibung(daten):
    berichte = []
    if daten["ort"].lower() == "remote" and not daten["homeoffice"]:
        berichte.append("❌ Widerspruch: 'Remote' gewählt, aber Homeoffice ist deaktiviert!")
    gewaehlte_stufe = next(
        (s for s in st.session_state.stufen_liste if s["name"].lower() == daten["level"].lower()),
        None
    )
    if gewaehlte_stufe and daten.get("erfahrung", 0) > 0:
        if daten["erfahrung"] < gewaehlte_stufe["erfahrung_min"]:
            berichte.append(f"❌ Erfahrung ({daten['erfahrung']} Jahre) zu wenig für {gewaehlte_stufe['name']} (Min: {gewaehlte_stufe['erfahrung_min']} Jahre).")
        if daten["erfahrung"] > gewaehlte_stufe["erfahrung_max"]:
            berichte.append(f"❌ Erfahrung ({daten['erfahrung']} Jahre) zu viel für {gewaehlte_stufe['name']} (Max: {gewaehlte_stufe['erfahrung_max']} Jahre).")
    return berichte

# --- WEB-OBERFLÄCHE ---
st.set_page_config(page_title="Recruitment-Check", page_icon="🏢", layout="wide")

# --- SIDEBAR ---
einstellungen = lade_hr_einstellungen()

if "stufen_liste" not in st.session_state:
    st.session_state.stufen_liste = einstellungen.get("stufen_liste", [
        {"name": "Junior", "erfahrung_min": 0, "erfahrung_max": 2, "gehalt_min": 35000, "gehalt_max": 50000},
        {"name": "Professional", "erfahrung_min": 2, "erfahrung_max": 5, "gehalt_min": 50000, "gehalt_max": 70000},
        {"name": "Senior", "erfahrung_min": 5, "erfahrung_max": 10, "gehalt_min": 70000, "gehalt_max": 90000},
        {"name": "Lead", "erfahrung_min": 10, "erfahrung_max": 99, "gehalt_min": 90000, "gehalt_max": 150000},
    ])

with st.sidebar:
    st.header("⚙️ HR-Einstellungen")
    st.caption("Diese Einstellungen werden von HR konfiguriert.")
    st.divider()
    FIRMEN_NAME = st.text_input("Firmenname", value=einstellungen["firmen_name"])
    ANREDE = st.selectbox("Anrede", ["Du", "Sie"],
        index=["Du", "Sie"].index(einstellungen["anrede"]))
    GENDER_STIL = st.selectbox("Gender-Stil", [
        "Gender-Sternchen (z.B. Mitarbeiter*innen)",
        "Doppelpunkt (z.B. Mitarbeiter:innen)",
        "Ausschreiben (z.B. Mitarbeiterinnen und Mitarbeiter)",
        "Kein Gendern"
    ], index=[
        "Gender-Sternchen (z.B. Mitarbeiter*innen)",
        "Doppelpunkt (z.B. Mitarbeiter:innen)",
        "Ausschreiben (z.B. Mitarbeiterinnen und Mitarbeiter)",
        "Kein Gendern"
    ].index(einstellungen["gender_stil"]))
    HIGHLIGHT = st.text_input("Highlight", value=einstellungen["highlight"])
    st.divider()
    st.subheader("📊 Erfahrungsstufen")
    st.caption("Stufen mit Erfahrungsjahren definieren")
    ERFAHRUNGSSTUFEN = []
    for i, stufe in enumerate(st.session_state.stufen_liste):
        with st.expander(f"📌 {stufe['name']}"):
            stufe["name"] = st.text_input("Stufenname", value=stufe["name"], key=f"stufe_name_{i}")
            col1, col2 = st.columns(2)
            with col1:
                stufe["erfahrung_min"] = st.number_input("Erfahrung Min (Jahre)", value=stufe["erfahrung_min"], key=f"erf_min_{i}")
                stufe["gehalt_min"] = st.number_input("Gehalt Min (€)", value=stufe["gehalt_min"], step=1000, key=f"geh_min_{i}")
            with col2:
                stufe["erfahrung_max"] = st.number_input("Erfahrung Max (Jahre)", value=stufe["erfahrung_max"], key=f"erf_max_{i}")
                stufe["gehalt_max"] = st.number_input("Gehalt Max (€)", value=stufe["gehalt_max"], step=1000, key=f"geh_max_{i}")
            if st.button("🗑️ Stufe löschen", key=f"delete_stufe_{i}"):
                st.session_state.stufen_liste.pop(i)
                st.rerun()
        ERFAHRUNGSSTUFEN.append(stufe["name"])
    if st.button("➕ Stufe hinzufügen"):
        st.session_state.stufen_liste.append({
            "name": "Neue Stufe",
            "erfahrung_min": 0,
            "erfahrung_max": 5,
            "gehalt_min": 40000,
            "gehalt_max": 60000
        })
        st.rerun()
    st.divider()
    if st.button("💾 Einstellungen speichern"):
        speichere_hr_einstellungen({
            "firmen_name": FIRMEN_NAME,
            "anrede": ANREDE,
            "gender_stil": GENDER_STIL,
            "highlight": HIGHLIGHT,
            "stufen_liste": st.session_state.stufen_liste
        })
        st.success("✅ Einstellungen gespeichert!")
        st.rerun()

# --- HAUPT-BEREICH ---
col_titel, col_reset = st.columns([5, 1])
with col_titel:
    st.title("Recruitment-Check Tool 🚀")
with col_reset:
    st.write("")
    st.write("")
    if st.button("🔄 Neu"):
        trigger = st.session_state.get("reset_trigger", 0) + 1
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.session_state.reset_trigger = trigger
        st.session_state.aktueller_schritt = 1
        st.rerun()

zeige_schritt_anzeige(st.session_state.aktueller_schritt)
st.divider()

# ============================================================
# SCHRITT 1: DATEN EINGEBEN
# ============================================================
if st.session_state.aktueller_schritt == 1:
    st.subheader("📝 Schritt 1: Stelle beschreiben")

    tab1, tab2 = st.tabs(["📝 Formular ausfüllen", "🤖 Agent – Stelle beschreiben"])

    # --- TAB 1: FORMULAR ---
    with tab1:
        col1, col2 = st.columns(2)
        with col1:
            job_titel = st.text_input(
                "Titel der Stelle",
                placeholder="z.B. IT Consultant",
                key=f"job_titel_{st.session_state.reset_trigger}")
            level = st.selectbox(
                "Erfahrungslevel", ERFAHRUNGSSTUFEN,
                key=f"level_{st.session_state.reset_trigger}")
        with col2:
            erfahrung = st.number_input(
                "Berufserfahrung (Jahre)", min_value=0, value=0,
                key=f"erfahrung_{st.session_state.reset_trigger}")
            ort = st.selectbox(
                "Arbeitsmodell", ["On-site", "Remote", "Hybrid"],
                key=f"ort_{st.session_state.reset_trigger}")
            homeoffice = st.checkbox(
                "Homeoffice-Option verfügbar?", value=True,
                key=f"homeoffice_{st.session_state.reset_trigger}")

        aufgaben = st.text_area(
            "Was sind die Hauptaufgaben?",
            placeholder="z.B. - Einkauf von IT-Hardware\n- Verhandlung mit Lieferanten",
            key=f"aufgaben_{st.session_state.reset_trigger}")

        st.divider()
        if st.button("▶️ Anzeige generieren →", type="primary", key="formular_generieren"):
            if not job_titel:
                st.warning("⚠️ Bitte gib mindestens einen Jobtitel ein.")
            else:
                with st.spinner("✍️ Anzeige wird erstellt..."):
                    try:
                        anzeige, fehler = generiere_anzeige(job_titel, level, erfahrung, ort, homeoffice, aufgaben)
                        if fehler:
                            st.error("Folgende Probleme wurden gefunden:")
                            for f in fehler:
                                st.write(f)
                        else:
                            st.session_state.anzeige = anzeige
                            st.session_state.historie.append({
                                "version": len(st.session_state.historie) + 1,
                                "inhalt": st.session_state.anzeige,
                                "typ": "Erste Version",
                                "datum": datetime.now().strftime("%d.%m.%Y %H:%M"),
                                "bearbeiter": "KI"
                            })
                            st.session_state.aktueller_schritt = 2
                            st.rerun()
                    except Exception as e:
                        st.error(f"Fehler: {e}")

    # --- TAB 2: AGENT ---
    with tab2:
        st.markdown("**Beschreibe einfach was du suchst – der Agent erstellt die Anzeige direkt für dich.**")

        if st.session_state.agent_rückfragen is None:
            freitext = st.text_area(
                "Was suchst du?",
                placeholder="z.B. 'Ich suche jemanden der unsere Kunden betreut, gut kommunizieren kann und Erfahrung im IT-Bereich hat'",
                height=150,
                key="agent_freitext_input"
            )
            if st.button("🤖 Agent starten"):
                if freitext:
                    with st.spinner("🤖 Agent analysiert deine Beschreibung..."):
                        analyse_prompt = f"""
                        Du bist ein HR-Experte für das Unternehmen "{FIRMEN_NAME}".
                        Ein Teamleiter von {FIRMEN_NAME} hat folgende Stelle beschrieben:
                        "{freitext}"
                        
                        WICHTIG:
                        - Erfinde KEINE Branche oder Unternehmensinfos
                        - Nutze NUR was der Teamleiter explizit beschrieben hat
                        - Falls Infos fehlen, stelle eine Rückfrage – erfinde nichts
                        
                        Folgende Informationen werden benötigt:
                        - Jobtitel (professionelle Bezeichnung)
                        - Erfahrungslevel (aus: {', '.join(ERFAHRUNGSSTUFEN)})
                        - Arbeitsmodell (On-site, Remote oder Hybrid)
                        - Hauptaufgaben (mindestens 3 konkrete Aufgaben)
                        - Berufserfahrung in Jahren
                        
                        Antworte NUR in diesem JSON Format ohne zusätzlichen Text:
                        {{
                            "verstanden": {{
                                "jobtitel": "...",
                                "level": "...",
                                "arbeitsmodell": "...",
                                "aufgaben": "...",
                                "erfahrung": 0
                            }},
                            "rückfragen": [
                                "Frage 1?",
                                "Frage 2?"
                            ]
                        }}
                        """
                        try:
                            rohtext = gemini(analyse_prompt).strip()
                            if rohtext.startswith("```"):
                                rohtext = rohtext.split("```")[1]
                                if rohtext.startswith("json"):
                                    rohtext = rohtext[4:]
                            result = json.loads(rohtext)
                            st.session_state.agent_rückfragen = result
                            st.session_state.agent_freitext = freitext
                            st.rerun()
                        except Exception as e:
                            st.error(f"Agent konnte den Text nicht analysieren: {e}")
                else:
                    st.info("Bitte beschreibe zuerst was du suchst.")

        else:
            verstanden = st.session_state.agent_rückfragen.get("verstanden", {})
            rückfragen = st.session_state.agent_rückfragen.get("rückfragen", [])

            if any(v for v in verstanden.values() if v):
                st.success("✅ Das habe ich aus deiner Beschreibung verstanden:")
                labels = {
                    "jobtitel": "Jobtitel",
                    "level": "Level",
                    "arbeitsmodell": "Arbeitsmodell",
                    "aufgaben": "Aufgaben",
                    "erfahrung": "Erfahrung"
                }
                for key, value in verstanden.items():
                    if value:
                        st.markdown(f"- **{labels.get(key, key)}:** {value}")

            if rückfragen:
                st.divider()
                st.markdown("**❓ Noch ein paar Fragen:**")
                for frage in rückfragen:
                    st.markdown(f"- {frage}")
                antworten = st.text_area(
                    "Deine Antworten:",
                    placeholder="Beantworte die Fragen oben...",
                    height=120,
                    key="agent_antworten_input"
                )
            else:
                antworten = ""
                st.info("✅ Ich habe alle Infos – klicke auf 'Anzeige generieren'!")

            st.divider()
            col1, col2 = st.columns(2)
            with col1:
                if st.button("▶️ Anzeige generieren →", type="primary", key="agent_generieren"):
                    with st.spinner("🤖 Agent erstellt die Anzeige..."):
                        try:
                            formular_prompt = f"""
                            Du bist ein HR-Experte für {FIRMEN_NAME}.
                            Ursprüngliche Beschreibung: "{st.session_state.agent_freitext}"
                            Was ich verstanden hatte: {json.dumps(verstanden, ensure_ascii=False)}
                            Antworten auf Rückfragen: "{antworten}"
                            
                            Erstelle jetzt ein vollständiges Formular.
                            Antworte NUR in diesem JSON Format ohne zusätzlichen Text:
                            {{
                                "jobtitel": "...",
                                "level": "...",
                                "arbeitsmodell": "...",
                                "aufgaben": "...",
                                "erfahrung": 0
                            }}
                            
                            Level muss einer dieser Werte sein: {', '.join(ERFAHRUNGSSTUFEN)}
                            Arbeitsmodell muss einer dieser Werte sein: On-site, Remote, Hybrid
                            Erfahrung ist eine Zahl (Jahre)
                            """
                            rohtext = gemini(formular_prompt).strip()
                            if rohtext.startswith("```"):
                                rohtext = rohtext.split("```")[1]
                                if rohtext.startswith("json"):
                                    rohtext = rohtext[4:]
                            formular = json.loads(rohtext)

                            anzeige, fehler = generiere_anzeige(
                                formular.get("jobtitel", ""),
                                formular.get("level", ERFAHRUNGSSTUFEN[0]),
                                int(formular.get("erfahrung", 0)),
                                formular.get("arbeitsmodell", "On-site"),
                                True,
                                formular.get("aufgaben", "")
                            )
                            if fehler:
                                for f in fehler:
                                    st.warning(f)
                            else:
                                st.session_state.anzeige = anzeige
                                st.session_state.job_titel_gespeichert = formular.get("jobtitel", "")
                                st.session_state.historie.append({
                                    "version": len(st.session_state.historie) + 1,
                                    "inhalt": st.session_state.anzeige,
                                    "typ": "Erste Version (via Agent)",
                                    "datum": datetime.now().strftime("%d.%m.%Y %H:%M"),
                                    "bearbeiter": "KI"
                                })
                                st.session_state.aktueller_schritt = 2
                                st.rerun()
                        except Exception as e:
                            st.error(f"Fehler: {e}")
            with col2:
                if st.button("🔄 Neu starten", key="agent_restart"):
                    st.session_state.agent_rückfragen = None
                    st.session_state.agent_freitext = ""
                    st.session_state.agent_antworten = ""
                    st.rerun()

# ============================================================
# SCHRITT 2: ANZEIGE PRÜFEN & FEEDBACK
# ============================================================
elif st.session_state.aktueller_schritt == 2:
    st.subheader("✍️ Schritt 2: Anzeige prüfen & anpassen")
    st.markdown(st.session_state.anzeige)
    st.divider()

    st.subheader("💬 Feedback vom Teamleiter")
    feedback = st.text_area(
        "Was soll geändert werden?",
        placeholder="z.B. 'Nimm das Studium raus' oder 'Füge Team Events hinzu'..."
    )

    col1, col2 = st.columns(2)
    with col1:
        if st.button("🔄 Änderungen übernehmen"):
            if feedback:
                refinement_prompt = f"""
                Du hast bereits diese Anzeige erstellt:
                {st.session_state.anzeige}
                
                Der Teamleiter hat nun folgendes Feedback gegeben:
                "{feedback}"
                
                Bitte überarbeite NUR den Inhalt der Anzeige.
                Behalte die Struktur und alle Überschriften exakt bei.
                WICHTIG: Das Feedback des Teamleiters steht ÜBER dem Wissen aus dem Archiv.
                """
                with st.spinner("✍️ Anzeige wird angepasst..."):
                    try:
                        st.session_state.anzeige = formatiere_ueberschriften(gemini(refinement_prompt))
                        st.session_state.historie.append({
                            "version": len(st.session_state.historie) + 1,
                            "inhalt": st.session_state.anzeige,
                            "typ": f"Feedback: {feedback[:50]}",
                            "datum": datetime.now().strftime("%d.%m.%Y %H:%M"),
                            "bearbeiter": "Teamleiter"
                        })
                        st.rerun()
                    except Exception as e:
                        st.error(f"Fehler: {e}")
            else:
                st.info("Bitte gib erst ein Feedback ein.")
    with col2:
        if st.button("📤 Anzeige an HR senden →", type="primary"):
            with st.spinner("⚖️ Finale Prüfung läuft..."):
                st.session_state.anzeige = fuehre_compliance_loop_durch(st.session_state.anzeige)
            st.session_state.hr_status = "wartend"
            st.session_state.aktueller_schritt = 3
            st.rerun()

    if st.session_state.historie:
        st.divider()
        st.subheader("📋 Versions-Historie")
        for i, version in enumerate(reversed(st.session_state.historie)):
            idx = len(st.session_state.historie) - 1 - i
            with st.expander(
                f"v{version['version']} | {version['datum']} | {version['bearbeiter']} | {version['typ']}",
                expanded=False
            ):
                if idx > 0:
                    st.caption("🟢 Hinzugefügt   🔴 Entfernt")
                    alt_zeilen = st.session_state.historie[idx - 1]["inhalt"].split("\n")
                    neu_zeilen = version["inhalt"].split("\n")
                    for neu_zeile in neu_zeilen:
                        if not neu_zeile.strip():
                            st.markdown("")
                            continue
                        beste_alte = None
                        beste_ratio = 0
                        for alt_zeile in alt_zeilen:
                            ratio = difflib.SequenceMatcher(None, alt_zeile, neu_zeile).ratio()
                            if ratio > beste_ratio:
                                beste_ratio = ratio
                                beste_alte = alt_zeile
                        if beste_ratio > 0.8:
                            diff_zeile = zeige_diff_wortweise(beste_alte, neu_zeile)
                            if neu_zeile.startswith("###") or neu_zeile.startswith("####"):
                                st.markdown(f"<div style='font-size:16px;font-weight:bold;margin-top:12px'>{diff_zeile}</div>", unsafe_allow_html=True)
                            elif neu_zeile.startswith("- "):
                                st.markdown(f"<div style='margin-left:16px'>• {diff_zeile}</div>", unsafe_allow_html=True)
                            else:
                                st.markdown(f"<div style='font-size:14px'>{diff_zeile}</div>", unsafe_allow_html=True)
                        elif beste_ratio < 0.3:
                            if neu_zeile.startswith("###") or neu_zeile.startswith("####"):
                                st.markdown(f"<div style='font-size:16px;font-weight:bold;margin-top:12px;background-color:#00cc4422;border-left:3px solid #00cc44;padding:4px 8px'>{neu_zeile.lstrip('#').strip()}</div>", unsafe_allow_html=True)
                            elif neu_zeile.startswith("- "):
                                st.markdown(f"<div style='margin-left:16px;background-color:#00cc4422;border-left:3px solid #00cc44;padding:4px 8px'>• {neu_zeile[2:]}</div>", unsafe_allow_html=True)
                            else:
                                st.markdown(f"<div style='font-size:14px;background-color:#00cc4422;border-left:3px solid #00cc44;padding:4px 8px'>{neu_zeile}</div>", unsafe_allow_html=True)
                        else:
                            diff_zeile = zeige_diff_wortweise(beste_alte, neu_zeile)
                            st.markdown(f"<div style='font-size:14px'>{diff_zeile}</div>", unsafe_allow_html=True)
                    for alt_zeile in alt_zeilen:
                        if not alt_zeile.strip():
                            continue
                        in_neu = any(
                            difflib.SequenceMatcher(None, alt_zeile, n).ratio() > 0.8
                            for n in neu_zeilen
                        )
                        if not in_neu:
                            st.markdown(f"<div style='font-size:14px;background-color:#ff444422;border-left:3px solid #ff4444;padding:4px 8px;text-decoration:line-through'>{alt_zeile.strip()}</div>", unsafe_allow_html=True)
                else:
                    st.caption("Erste Version")
                    st.markdown(version["inhalt"])

# ============================================================
# SCHRITT 3: HR FREIGABE
# ============================================================
elif st.session_state.aktueller_schritt == 3:
    st.subheader("📤 Schritt 3: HR Freigabe")

    if st.session_state.hr_status == "wartend":
        st.info("⏳ Wartet auf HR-Freigabe...")
        st.markdown("---")
        st.markdown("**HR-Bereich – Finale Anzeige:**")
        st.markdown(st.session_state.anzeige)
        st.divider()

        hr_kommentar = st.text_area(
            "Kommentar für Teamleiter (optional):",
            placeholder="z.B. 'Bitte Anforderungen anpassen...'"
        )
        if hr_kommentar:
            st.session_state.hr_kommentar = hr_kommentar

        col1, col2 = st.columns(2)
        with col1:
            if st.button("✅ Freigeben", type="primary"):
                st.session_state.hr_status = "freigegeben"
                st.session_state.aktueller_schritt = 4
                st.rerun()
        with col2:
            if st.button("❌ Ablehnen"):
                st.session_state.hr_status = "abgelehnt"
                st.rerun()

    elif st.session_state.hr_status == "abgelehnt":
        st.error("❌ Anzeige wurde von HR abgelehnt!")
        if st.session_state.hr_kommentar:
            st.info(f"💬 Begründung HR: {st.session_state.hr_kommentar}")
        st.divider()
        if st.button("🔄 Zurück zur Anzeige überarbeiten"):
            st.session_state.hr_status = None
            st.session_state.aktueller_schritt = 2
            st.rerun()

# ============================================================
# SCHRITT 4: EXPORT
# ============================================================
elif st.session_state.aktueller_schritt == 4:
    st.subheader("📥 Schritt 4: Export")
    st.success("✅ Anzeige wurde von HR freigegeben!")
    if st.session_state.hr_kommentar:
        st.info(f"💬 Kommentar HR: {st.session_state.hr_kommentar}")

    st.divider()
    st.markdown(st.session_state.anzeige)
    st.divider()

    dateiname_basis = st.session_state.job_titel_gespeichert.replace(' ', '_')
    col1, col2 = st.columns(2)
    with col1:
        word_buffer = erstelle_word_dokument(st.session_state.anzeige)
        st.download_button(
            label="📄 Als Word herunterladen",
            data=word_buffer,
            file_name=f"stellenanzeige_{dateiname_basis}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="btn_download_word"
        )
    with col2:
        pdf_buffer = erstelle_pdf_dokument(
            st.session_state.anzeige,
            st.session_state.job_titel_gespeichert
        )
        st.download_button(
            label="📑 Als PDF herunterladen",
            data=pdf_buffer,
            file_name=f"stellenanzeige_{dateiname_basis}.pdf",
            mime="application/pdf",
            key="btn_download_pdf"
        )